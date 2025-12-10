import os
import sys
import glob
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.shared import Inches        # >>> NEW: for adding images to docx
import matplotlib.pyplot as plt
from sklearn.decomposition import PCA
from sklearn.decomposition import FastICA
from sklearn.preprocessing import MinMaxScaler
from sklearn.decomposition import FactorAnalysis
from sklearn.model_selection import train_test_split
from sklearn.feature_selection import mutual_info_classif, chi2, f_classif
from sklearn.metrics import f1_score, classification_report, confusion_matrix

import tensorflow as tf
import time
import psutil                      # >>> NEW: for resource usage (memory)

# GLOBAL CONFIGURATION
start_time = time.time()

if len(sys.argv) < 2:
    print("Usage: python CIC_IDS_Training.py <data_path>")
    sys.exit(1)

FUSION_LIST = ['pca', 'ica', 'fa']
MODEL_LIST = ['LSTM', 'ANN', 'GRU', 'CNN']
RANDOM_STATE = 42
TRAING_EPOCHS = 50
BATCH_SIZE = 512
SPLIT_SIZE = 0.2
RANGE_LIMIT = 75       # Upper limit for k in feature selection
MI_SAMPLE_SIZE = 200000  # >>> NEW: sample size for feature scoring to avoid OOM

data_path = sys.argv[1]

# Folders for results
OUTPUT_DIR = "model_results"
PLOT_DIR = "plots"
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(PLOT_DIR, exist_ok=True)

# GPU SETUP
gpus = tf.config.list_physical_devices('GPU')
if gpus:
    try:
        for gpu in gpus:
            tf.config.experimental.set_memory_growth(gpu, True)
        print("✅ Memory growth enabled on GPU.")
    except RuntimeError as e:
        print("⚠️ Could not set memory growth:", e)
else:
    print("⚠️ No GPU found. Running on CPU.")

from tensorflow.keras.layers import Conv1D, MaxPooling1D, Flatten, GRU
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import LSTM, Dense, Dropout, Bidirectional, BatchNormalization
from tensorflow.keras.callbacks import EarlyStopping


def load_data_from_directory(data_path):
    df = pd.read_csv(data_path)
    print("Dataset loaded sucessfully")
    return df


def preprocess_data(df):
    df.replace([np.inf, -np.inf], np.nan, inplace=True)
    df.dropna(inplace=True)
    print(f"Data shape after dropping NaN and infinite values: {df.shape}")

    X_fin_capped = df.drop(columns=["target"])
    y_encoded = df["target"]

    print("Class distribution:")
    print(y_encoded.value_counts())

    print('Preprocessing done.')
    return X_fin_capped, y_encoded


def split_sub_data(X_fin_capped, y_encoded, selected, test_size=0.2, random_state=42):
    print("Selected Features:", selected)
    print("New data created .")

    X_selected = X_fin_capped[selected]
    print("Feature selection done.")

    print("Original shape:", X_fin_capped.shape)
    print("Selected shape:", X_selected.shape)

    X_train, X_test, y_train, y_test = train_test_split(
        X_selected, y_encoded, test_size=test_size, random_state=random_state
    )
    print("Train test split done.")

    y_train = np.array(y_train)
    y_test = np.array(y_test)

    train_total = len(y_train)
    train_class0 = np.sum(y_train == 0)
    train_class1 = np.sum(y_train == 1)

    test_total = len(y_test)
    test_class0 = np.sum(y_test == 0)
    test_class1 = np.sum(y_test == 1)

    print("\n====== Class Distribution After Split ======")
    print(f"Training Total  : {train_total}")
    print(f"Training Class 0: {train_class0}")
    print(f"Training Class 1: {train_class1}")
    print("-------------------------------------------")
    print(f"Testing Total   : {test_total}")
    print(f"Testing Class 0 : {test_class0}")
    print(f"Testing Class 1 : {test_class1}")
    print("===========================================\n")

    return X_train, X_test, y_train, y_test


def calculate_Score(X, y):
    print("calculated mi score.")
    mi_scores = mutual_info_classif(X, y, random_state=42)
    print("Sucessfully calculated mi score.")

    print("calculated chi score.")
    X_chi = X.copy()
    X_chi[X_chi < 0] = 0
    chi_scores, _ = chi2(X_chi, y)
    print("Sucessfully calculated chi score.")

    print("calculated f score.")
    f_scores, _ = f_classif(X, y)
    print("Sucessfully calculated f score.")

    return mi_scores, chi_scores, f_scores


def hybrid_feature_selection_pca(mi_scores, chi_scores, f_scores, feature_names, k):
    print("Starting PCA-based feature fusion...")

    scores = np.vstack([mi_scores, chi_scores, f_scores]).T

    scaler = MinMaxScaler()
    scores_scaled = scaler.fit_transform(scores)

    pca = PCA(n_components=1)
    meta_score = pca.fit_transform(scores_scaled).ravel()

    feature_df = pd.DataFrame({
        'Feature': feature_names,
        'MI_Score': mi_scores,
        'Chi2_Score': chi_scores,
        'F_Score': f_scores,
        'Meta_Score': meta_score
    }).sort_values(by='Meta_Score', ascending=False)

    selected_features = feature_df.head(k)['Feature'].tolist()

    print("✅ PCA fusion completed.")
    print("Top", k, "Selected Features:", selected_features)
    print("\nPCA Component Weights (importance of MI, Chi², F):")
    print(pd.Series(pca.components_[0], index=['MI', 'Chi2', 'F']))

    return selected_features, feature_df


def hybrid_feature_selection_ica(mi_scores, chi_scores, f_scores, feature_names, k):
    print("Starting ICA-based feature fusion...")

    scores = np.vstack([mi_scores, chi_scores, f_scores]).T

    scaler = MinMaxScaler()
    scores_scaled = scaler.fit_transform(scores)

    ica = FastICA(n_components=1, random_state=42, whiten='unit-variance')
    meta_score = ica.fit_transform(scores_scaled).ravel()

    simple_average = np.mean(scores_scaled, axis=1)
    correlation = np.corrcoef(meta_score, simple_average)[0, 1]

    if correlation < 0:
        print("🔄 ICA sign flip detected. Inverting scores to match reality...")
        meta_score = -meta_score

    feature_df = pd.DataFrame({
        'Feature': feature_names,
        'MI_Score': mi_scores,
        'Chi2_Score': chi_scores,
        'F_Score': f_scores,
        'ICA_Meta_Score': meta_score
    }).sort_values(by='ICA_Meta_Score', ascending=False)

    selected_features = feature_df.head(k)['Feature'].tolist()

    print("✅ ICA fusion completed.")
    print(f"Top {k} Selected Features: {selected_features}")

    return selected_features, feature_df


def hybrid_feature_selection_fa(mi_scores, chi_scores, f_scores, feature_names, k):
    print("Starting Factor Analysis (FA) feature fusion...")

    scores = np.vstack([mi_scores, chi_scores, f_scores]).T

    scaler = MinMaxScaler()
    scores_scaled = scaler.fit_transform(scores)

    fa = FactorAnalysis(n_components=1, random_state=42)
    meta_score = fa.fit_transform(scores_scaled).ravel()

    feature_df = pd.DataFrame({
        'Feature': feature_names,
        'MI_Score': mi_scores,
        'Chi2_Score': chi_scores,
        'F_Score': f_scores,
        'Meta_Score': meta_score
    }).sort_values(by='Meta_Score', ascending=False)

    selected_features = feature_df.head(k)['Feature'].tolist()

    print("✅ FA fusion completed.")
    print(f"Top {k} Selected Features: {selected_features}")

    print("\nFactor Loadings (Correlation with MI, Chi2, F):")
    print(pd.Series(fa.components_[0], index=['MI', 'Chi2', 'F']))

    return selected_features, feature_df


def feature_selection(X_fin_capped, mi_scores, chi_scores, f_scores, k=20, method='pca'):
    if method == 'pca':
        selected_features, score_table = hybrid_feature_selection_pca(
            mi_scores, chi_scores, f_scores,
            feature_names=X_fin_capped.columns, k=k
        )
    elif method == 'ica':
        selected_features, score_table = hybrid_feature_selection_ica(
            mi_scores, chi_scores, f_scores,
            feature_names=X_fin_capped.columns, k=k
        )
    elif method == 'fa':
        selected_features, score_table = hybrid_feature_selection_fa(
            mi_scores, chi_scores, f_scores,
            feature_names=X_fin_capped.columns, k=k
        )
    else:
        raise ValueError("Invalid method. Choose from 'pca', 'ica', or 'fa'.")

    return selected_features, score_table


# ===========================================
#  MODULE: UNIVERSAL TRAINING PIPELINE
# ===========================================
def run_single_model(model_name, X_fin_capped, y_encoded, feature_method,
                     mi_scores, chi_scores, f_scores):

    results = []
    print(f"\n============== {model_name} TRAINING STARTED ({feature_method}) ==============")

    process = psutil.Process(os.getpid())  # >>> NEW: track memory
    best_test_f1 = -1
    best_cm = None
    best_k_for_cm = None

    for k in range(5, RANGE_LIMIT + 1, 5):

        print(f"\n===== Testing k = {k} =====")
        print("Feature selection start:")

        k_start_time = time.time()  # >>> NEW: timing per k

        selected, value_table = feature_selection(
            X_fin_capped=X_fin_capped,
            mi_scores=mi_scores,
            chi_scores=chi_scores,
            f_scores=f_scores,
            k=k,
            method=feature_method
        )

        X_train, X_test, y_train, y_test = split_sub_data(
            X_fin_capped=X_fin_capped,
            y_encoded=y_encoded,
            selected=selected,
            test_size=SPLIT_SIZE,
            random_state=RANDOM_STATE
        )

        if model_name == "LSTM":
            model = Sequential()
            model.add(LSTM(128, input_shape=(1, X_train.shape[1]), return_sequences=True))
            model.add(BatchNormalization())
            model.add(Dropout(0.3))
            model.add(LSTM(64, return_sequences=True))
            model.add(Dropout(0.3))
            model.add(LSTM(64, return_sequences=False))
            model.add(BatchNormalization())
            model.add(Dropout(0.3))
            model.add(Dense(64, activation='relu'))
            model.add(Dropout(0.3))
            model.add(Dense(32, activation='relu'))
            model.add(Dense(1, activation='sigmoid'))

            X_train_m = X_train.values.reshape((X_train.shape[0], 1, X_train.shape[1]))
            X_test_m = X_test.values.reshape((X_test.shape[0], 1, X_test.shape[1]))

        elif model_name == "ANN":
            model = Sequential()
            model.add(tf.keras.Input(shape=(X_train.shape[1],)))
            model.add(Dense(128, activation='relu'))
            model.add(BatchNormalization())
            model.add(Dropout(0.3))
            model.add(Dense(64, activation='relu'))
            model.add(BatchNormalization())
            model.add(Dropout(0.3))
            model.add(Dense(32, activation='relu'))
            model.add(Dropout(0.3))
            model.add(Dense(1, activation='sigmoid'))
            X_train_m, X_test_m = X_train, X_test

        elif model_name == "GRU":
            model = Sequential()
            model.add(GRU(128, input_shape=(1, X_train.shape[1]), return_sequences=True))
            model.add(BatchNormalization())
            model.add(Dropout(0.3))
            model.add(GRU(64, return_sequences=False))
            model.add(BatchNormalization())
            model.add(Dropout(0.3))
            model.add(Dense(64, activation='relu'))
            model.add(Dropout(0.3))
            model.add(Dense(32, activation='relu'))
            model.add(Dense(1, activation='sigmoid'))
            X_train_m = X_train.values.reshape((X_train.shape[0], 1, X_train.shape[1]))
            X_test_m = X_test.values.reshape((X_test.shape[0], 1, X_test.shape[1]))

        elif model_name == "CNN":
            model = Sequential()
            model.add(tf.keras.Input(shape=(X_train.shape[1], 1)))

            model.add(Conv1D(filters=64, kernel_size=min(2, X_train.shape[1]),
                             activation='relu', padding='same'))
            model.add(BatchNormalization())
            if X_train.shape[1] >= 4:
                model.add(MaxPooling1D(pool_size=2))
            model.add(Dropout(0.3))

            model.add(Conv1D(filters=128, kernel_size=min(2, X_train.shape[1]),
                             activation='relu', padding='same'))
            model.add(BatchNormalization())
            if X_train.shape[1] > 8:
                model.add(MaxPooling1D(pool_size=2))
            model.add(Dropout(0.3))

            model.add(Flatten())
            model.add(Dense(64, activation='relu'))
            model.add(Dropout(0.3))
            model.add(Dense(32, activation='relu'))
            model.add(Dense(1, activation='sigmoid'))

            X_train_m = X_train.values.reshape((X_train.shape[0], X_train.shape[1], 1))
            X_test_m = X_test.values.reshape((X_test.shape[0], X_test.shape[1], 1))

        model.compile(loss='binary_crossentropy', optimizer='adam', metrics=['accuracy'])

        tf.debugging.set_log_device_placement(True)
        tf.config.set_soft_device_placement(True)

        gpus = tf.config.list_physical_devices('GPU')
        if gpus:
            try:
                for gpu in gpus:
                    tf.config.experimental.set_memory_growth(gpu, True)
            except RuntimeError as e:
                print("Memory growth error:", e)

        print("Training started...")
        history = model.fit(
            X_train_m, y_train,
            validation_data=(X_test_m, y_test),
            epochs=TRAING_EPOCHS,
            batch_size=BATCH_SIZE,
            verbose=1
        )
        print("Training done.")

        train_loss, train_acc = model.evaluate(X_train_m, y_train, verbose=0)
        test_loss, test_acc = model.evaluate(X_test_m, y_test, verbose=0)

        y_train_pred = (model.predict(X_train_m) > 0.5).astype(int)
        y_test_pred = (model.predict(X_test_m) > 0.5).astype(int)

        train_f1 = f1_score(y_train, y_train_pred)
        test_f1 = f1_score(y_test, y_test_pred)

        k_time = time.time() - k_start_time           # >>> NEW: time per k
        mem_gb = process.memory_info().rss / (1024**3)  # >>> NEW: memory in GB

        results.append({
            "k": k,
            "train_acc": train_acc,
            "test_acc": test_acc,
            "train_f1": train_f1,
            "test_f1": test_f1,
            "time_sec": k_time,
            "mem_gb": mem_gb
        })

        print(f"Train Acc = {train_acc:.4f}, Test Acc = {test_acc:.4f}")
        print(f"Train F1  = {train_f1:.4f}, Test F1  = {test_f1:.4f}")
        print(f"Time for k={k}: {k_time:.2f} sec, Memory: {mem_gb:.2f} GB")

        # Track best confusion matrix for this model + feature method
        if test_f1 > best_test_f1:
            best_test_f1 = test_f1
            best_cm = confusion_matrix(y_test, y_test_pred)
            best_k_for_cm = k

    results_df = pd.DataFrame(results)
    best_row = results_df.loc[results_df['test_f1'].idxmax()]
    print(f"\n{model_name} ({feature_method}) : 🏆 Best k = {best_row['k']} | Test F1 = {best_row['test_f1']:.4f}")

    # >>> NEW: Save per-model F1 vs k plot
    plt.figure()
    plt.plot(results_df['k'], results_df['train_f1'], marker='o', label='Train F1')
    plt.plot(results_df['k'], results_df['test_f1'], marker='o', label='Test F1')
    plt.xlabel("Number of Features (k)")
    plt.ylabel("F1 Score")
    plt.title(f"{model_name} - {feature_method.upper()} - F1 vs k")
    plt.legend()
    plt.grid(True)
    f1_plot_path = os.path.join(PLOT_DIR, f"{feature_method}_{model_name}_F1_vs_k.png")
    plt.savefig(f1_plot_path, bbox_inches='tight')
    plt.close()

    # >>> NEW: Save per-model time vs k plot (resource usage)
    plt.figure()
    plt.plot(results_df['k'], results_df['time_sec'], marker='o')
    plt.xlabel("Number of Features (k)")
    plt.ylabel("Training Time per run (sec)")
    plt.title(f"{model_name} - {feature_method.upper()} - Time vs k")
    plt.grid(True)
    time_plot_path = os.path.join(PLOT_DIR, f"{feature_method}_{model_name}_Time_vs_k.png")
    plt.savefig(time_plot_path, bbox_inches='tight')
    plt.close()

    # >>> NEW: Save confusion matrix heatmap for best k
    if best_cm is not None:
        plt.figure()
        sns.heatmap(best_cm, annot=True, fmt='d', cmap='Blues')
        plt.xlabel("Predicted")
        plt.ylabel("True")
        plt.title(f"{model_name} - {feature_method.upper()} - Confusion Matrix (k={best_k_for_cm})")
        cm_plot_path = os.path.join(PLOT_DIR, f"{feature_method}_{model_name}_Confusion_Matrix.png")
        plt.savefig(cm_plot_path, bbox_inches='tight')
        plt.close()

    # Also save raw CSV of results
    csv_path = os.path.join(OUTPUT_DIR, f"{feature_method}_{model_name}_results.csv")
    results_df.to_csv(csv_path, index=False)

    return results_df, best_row


# ===========================================
#  MASTER FUNCTION TO RUN ALL MODELS
# ===========================================
def run_all_models(X_fin_capped, y_encoded, mi_scores, chi_scores, f_scores, feature_method="pca"):

    fm_start_time = time.time()  # >>> NEW: time per feature fusion

    print(f"\n########## RUNNING ALL MODELS WITH {feature_method.upper()} FEATURE FUSION ##########")

    document = Document()
    document.add_heading(f"{feature_method.upper()} FULL MODEL RESULTS", level=1)

    # >>> NEW: Configuration + basic info in DOCX
    document.add_paragraph(f"Dataset shape: {X_fin_capped.shape[0]} rows, {X_fin_capped.shape[1]} features")
    document.add_paragraph(f"Train/Test split: {SPLIT_SIZE}, Random state: {RANDOM_STATE}")
    document.add_paragraph(f"Epochs: {TRAING_EPOCHS}, Batch size: {BATCH_SIZE}")
    document.add_paragraph(f"Feature fusion method: {feature_method.upper()}")

    model_results_dict = {}  # >>> NEW: store dfs for overall plot

    models = ["LSTM", "ANN", "GRU", "CNN"]

    for model in models:

        df, best = run_single_model(
            model_name=model,
            X_fin_capped=X_fin_capped,
            y_encoded=y_encoded,
            feature_method=feature_method,
            mi_scores=mi_scores,
            chi_scores=chi_scores,
            f_scores=f_scores
        )

        model_results_dict[model] = df

        document.add_heading(f"{model} Results", level=2)

        # Create table with header row
        table = document.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light List Accent 1'

        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                value = row[col]
                row_cells[i].text = str(value)

        document.add_paragraph(
            f"\n🏆 Best k = {best['k']} | Test F1 = {best['test_f1']:.4f}"
        )

        # >>> NEW: Add plots into DOCX if exist
        f1_plot_path = os.path.join(PLOT_DIR, f"{feature_method}_{model}_F1_vs_k.png")
        time_plot_path = os.path.join(PLOT_DIR, f"{feature_method}_{model}_Time_vs_k.png")
        cm_plot_path = os.path.join(PLOT_DIR, f"{feature_method}_{model}_Confusion_Matrix.png")

        if os.path.exists(f1_plot_path):
            document.add_paragraph("F1 Score vs Number of Features:")
            document.add_picture(f1_plot_path, width=Inches(5))

        if os.path.exists(time_plot_path):
            document.add_paragraph("Training Time vs Number of Features:")
            document.add_picture(time_plot_path, width=Inches(5))

        if os.path.exists(cm_plot_path):
            document.add_paragraph("Confusion Matrix (Best k):")
            document.add_picture(cm_plot_path, width=Inches(4))

        document.add_page_break()

    # >>> NEW: Overall comparison plot (all models, same feature method)
    plt.figure()
    for model, df in model_results_dict.items():
        plt.plot(df['k'], df['test_f1'], marker='o', label=model)
    plt.xlabel("Number of Features (k)")
    plt.ylabel("Test F1 Score")
    plt.title(f"Test F1 vs k for all models ({feature_method.upper()} fusion)")
    plt.legend()
    plt.grid(True)
    overall_plot_path = os.path.join(PLOT_DIR, f"{feature_method}_ALL_MODELS_F1_vs_k.png")
    plt.savefig(overall_plot_path, bbox_inches='tight')
    plt.close()

    if os.path.exists(overall_plot_path):
        document.add_heading("Overall Model Comparison (Test F1 vs k)", level=2)
        document.add_picture(overall_plot_path, width=Inches(5))

    fm_end_time = time.time()
    total_min = (fm_end_time - fm_start_time) / 60.0
    document.add_paragraph(f"\nTotal time for {feature_method.upper()} fusion: {total_min:.2f} minutes")

    filename = os.path.join(OUTPUT_DIR, f"ALL_MODEL_RESULTS_{feature_method.upper()}.docx")
    document.save(filename)

    print(f"\n📄 ALL MODEL RESULTS SAVED SUCCESSFULLY AS: {filename}")
    print(f"⏱ Time for {feature_method.upper()} fusion: {total_min:.2f} minutes")


if __name__ == "__main__":
    print("Data path received:", data_path)
    df = load_data_from_directory(data_path)

    X_fin_capped, y_encoded = preprocess_data(df)

    # >>> NEW: Sample for MI/Chi2/F to avoid KILLED
    if X_fin_capped.shape[0] > MI_SAMPLE_SIZE:
        print(f"Using a random sample of {MI_SAMPLE_SIZE} rows for feature scoring to avoid OOM.")
        rng = np.random.RandomState(RANDOM_STATE)
        idx = rng.choice(X_fin_capped.shape[0], MI_SAMPLE_SIZE, replace=False)
        X_score = X_fin_capped.iloc[idx]
        y_score = y_encoded.iloc[idx]
    else:
        X_score, y_score = X_fin_capped, y_encoded

    mi_scores, chi_scores, f_scores = calculate_Score(X_score, y_score)

    fusion_list = ['pca', 'ica', 'fa']

    for method in fusion_list:
        run_all_models(
            X_fin_capped, y_encoded,
            mi_scores=mi_scores,
            chi_scores=chi_scores,
            f_scores=f_scores,
            feature_method=method
        )

    end_time = time.time()
    total_minutes = (end_time - start_time) / 60.0
    print(f"\n✅ All processes completed successfully.")
    print(f"⏱ TOTAL EXECUTION TIME: {total_minutes:.2f} minutes")
