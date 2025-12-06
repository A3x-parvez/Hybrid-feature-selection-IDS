#!/usr/bin/env python3
"""
Hybrid-feature-selection-IDS
Single-file, dimension-safe, GPU-safe pipeline that:
- Loads dataset (auto-detects 'target' column)
- Preprocesses (drops NaN/inf, numeric-only, MinMax scaling)
- Computes MI / Chi2 / F-score once per fusion method
- Performs hybrid fusion (PCA / ICA / FA)
- For k in range(5..60 step 5, clipped to feature count) trains LSTM / ANN / GRU / CNN
- Uses EarlyStopping (restore_best_weights) for faster runs
- Saves results per fusion method into a DOCX report with tables
"""

import os
import sys
import time
import numpy as np
import pandas as pd
from docx import Document
from sklearn.decomposition import PCA, FastICA, FactorAnalysis
from sklearn.preprocessing import MinMaxScaler
from sklearn.feature_selection import mutual_info_classif, chi2, f_classif
from sklearn.metrics import f1_score
import tensorflow as tf

# ---------------------------
# GPU / TF setup (must be before Keras imports)
# ---------------------------
gpus = tf.config.list_physical_devices('GPU')
if gpus:
    try:
        for gpu in gpus:
            tf.config.experimental.set_memory_growth(gpu, True)
        print("✅ GPU detected — memory growth enabled.")
    except RuntimeError as e:
        print("⚠️ Could not set memory growth:", e)
else:
    print("ℹ️ No GPU found. Running on CPU.")

# Now safe to import Keras layers
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import (
    LSTM, GRU, Dense, Dropout, BatchNormalization,
    Conv1D, MaxPooling1D, Flatten
)
from tensorflow.keras.callbacks import EarlyStopping

# ---------------------------
# Configuration (tweak here)
# ---------------------------
RANDOM_STATE = 42
TRAING_EPOCHS = 20          # per-run epochs (EarlyStopping may stop earlier)
BATCH_SIZE = 512
DEFAULT_SPLIT = 0.2
K_RANGE = list(range(5, 65, 5))   # candidate k values; will be clipped to available features
VERBOSE = 1                      # model.fit verbose (0,1,2)
DOCX_TABLE_STYLE = 'Light List Accent 1'  # MS Word style; may vary by Word theme

# Reproducibility
np.random.seed(RANDOM_STATE)
tf.random.set_seed(RANDOM_STATE)

# ---------------------------
# Utilities
# ---------------------------
def find_target_column(df):
    """Return the name of the target column if found among common names."""
    candidates = ["target", "label", "class", "y", "Target", "Label", "CLASS"]
    for c in candidates:
        if c in df.columns:
            return c
    # fallback: if last column looks binary (0/1), pick it
    last_col = df.columns[-1]
    vals = pd.unique(df[last_col])
    if set(vals).issubset({0,1}) and len(vals) <= 2:
        return last_col
    raise ValueError("Target column not found. Please include a 'target' column (binary 0/1) or one of: "
                     + ", ".join(candidates))

def safe_minmax_scale(df):
    """Scale numeric dataframe columns to [0,1] and return scaled df and scaler."""
    scaler = MinMaxScaler()
    scaled = scaler.fit_transform(df.values)
    return pd.DataFrame(scaled, columns=df.columns, index=df.index)

# ---------------------------
# Feature scoring & fusion
# ---------------------------
def calculate_Score(X, y):
    """Compute MI, Chi2 (on non-negative copy), and ANOVA F-scores."""
    print("🔢 Calculating feature-level scores (MI / Chi2 / F)...")
    mi_scores = mutual_info_classif(X, y, random_state=RANDOM_STATE)
    # chi2 requires non-negative values
    X_chi = X.copy()
    X_chi[X_chi < 0] = 0
    chi_scores, _ = chi2(X_chi, y)
    f_scores, _ = f_classif(X, y)
    print("✅ Scores calculated.")
    return mi_scores, chi_scores, f_scores

def hybrid_feature_selection_pca(mi_scores, chi_scores, f_scores, feature_names, k):
    scores = np.vstack([mi_scores, chi_scores, f_scores]).T
    scaler = MinMaxScaler()
    scores_scaled = scaler.fit_transform(scores)
    pca = PCA(n_components=1, random_state=RANDOM_STATE)
    meta_score = pca.fit_transform(scores_scaled).ravel()
    feature_df = pd.DataFrame({
        'Feature': feature_names,
        'MI_Score': mi_scores,
        'Chi2_Score': chi_scores,
        'F_Score': f_scores,
        'Meta_Score': meta_score
    }).sort_values(by='Meta_Score', ascending=False)
    selected = feature_df.head(k)['Feature'].tolist()
    return selected, feature_df

def hybrid_feature_selection_ica(mi_scores, chi_scores, f_scores, feature_names, k):
    scores = np.vstack([mi_scores, chi_scores, f_scores]).T
    scaler = MinMaxScaler()
    scores_scaled = scaler.fit_transform(scores)
    ica = FastICA(n_components=1, random_state=RANDOM_STATE, whiten='unit-variance')
    meta_score = ica.fit_transform(scores_scaled).ravel()
    # handle sign flip
    avg = np.mean(scores_scaled, axis=1)
    corr = np.corrcoef(meta_score, avg)[0,1]
    if corr < 0:
        meta_score = -meta_score
    feature_df = pd.DataFrame({
        'Feature': feature_names,
        'MI_Score': mi_scores,
        'Chi2_Score': chi_scores,
        'F_Score': f_scores,
        'ICA_Meta_Score': meta_score
    }).sort_values(by='ICA_Meta_Score', ascending=False)
    selected = feature_df.head(k)['Feature'].tolist()
    return selected, feature_df

def hybrid_feature_selection_fa(mi_scores, chi_scores, f_scores, feature_names, k):
    scores = np.vstack([mi_scores, chi_scores, f_scores]).T
    scaler = MinMaxScaler()
    scores_scaled = scaler.fit_transform(scores)
    fa = FactorAnalysis(n_components=1, random_state=RANDOM_STATE)
    meta_score = fa.fit_transform(scores_scaled).ravel()
    feature_df = pd.DataFrame({
        'Feature': feature_names,
        'MI_Score': mi_scores,
        'Chi2_Score': chi_scores,
        'F_Score': f_scores,
        'Meta_Score': meta_score
    }).sort_values(by='Meta_Score', ascending=False)
    selected = feature_df.head(k)['Feature'].tolist()
    return selected, feature_df

def feature_selection(X_fin_capped, mi_scores, chi_scores, f_scores, k=20, method='pca'):
    if k <= 0:
        raise ValueError("k must be positive")
    method = method.lower()
    if method == 'pca':
        return hybrid_feature_selection_pca(mi_scores, chi_scores, f_scores, feature_names=X_fin_capped.columns, k=k)
    elif method == 'ica':
        return hybrid_feature_selection_ica(mi_scores, chi_scores, f_scores, feature_names=X_fin_capped.columns, k=k)
    elif method == 'fa':
        return hybrid_feature_selection_fa(mi_scores, chi_scores, f_scores, feature_names=X_fin_capped.columns, k=k)
    else:
        raise ValueError("method must be one of 'pca','ica','fa'")

# ---------------------------
# Data split helper
# ---------------------------
from sklearn.model_selection import train_test_split
def split_sub_data(X_fin_capped, y_encoded, selected, test_size=DEFAULT_SPLIT, random_state=RANDOM_STATE):
    # Ensure selected are valid columns
    selected = [c for c in selected if c in X_fin_capped.columns]
    X_selected = X_fin_capped[selected]
    X_train, X_test, y_train, y_test = train_test_split(
        X_selected, y_encoded, test_size=test_size, random_state=random_state, stratify=y_encoded
    )
    # convert to numpy arrays for ML functions
    y_train = np.array(y_train)
    y_test = np.array(y_test)
    # print class distribution
    train_total = len(y_train)
    test_total = len(y_test)
    train_class0 = int(np.sum(y_train == 0))
    train_class1 = int(np.sum(y_train == 1))
    test_class0 = int(np.sum(y_test == 0))
    test_class1 = int(np.sum(y_test == 1))
    print(f"Train: total={train_total}, class0={train_class0}, class1={train_class1} | "
          f"Test: total={test_total}, class0={test_class0}, class1={test_class1}")
    return X_train, X_test, y_train, y_test

# ---------------------------
# Model runner (architectures kept unchanged)
# ---------------------------
def run_single_model(model_name, X_fin_capped, y_encoded, feature_method, mi_scores, chi_scores, f_scores):
    results = []
    n_features_available = X_fin_capped.shape[1]
    # compute effective k list clipped by available features
    effective_k = [k for k in K_RANGE if k <= n_features_available]
    if len(effective_k) == 0:
        # fallback to use min(5, n_features)
        effective_k = [min(5, n_features_available)]
    print(f"\n>>> Running {model_name} across k={effective_k} (method={feature_method})")
    for k in effective_k:
        print(f"\n--- k = {k} ---")
        selected, score_table = feature_selection(
            X_fin_capped=X_fin_capped, mi_scores=mi_scores, chi_scores=chi_scores, f_scores=f_scores, k=k, method=feature_method
        )
        X_train, X_test, y_train, y_test = split_sub_data(X_fin_capped, y_encoded, selected)

        # --- build model (exact same architecture & hyperparams as before) ---
        if model_name == "LSTM":
            model = Sequential()
            model.add(LSTM(128, input_shape=(1, X_train.shape[1]), return_sequences=True))
            model.add(BatchNormalization()); model.add(Dropout(0.3))
            model.add(LSTM(64, return_sequences=True)); model.add(Dropout(0.3))
            model.add(LSTM(64, return_sequences=False))
            model.add(BatchNormalization()); model.add(Dropout(0.3))
            model.add(Dense(64, activation='relu')); model.add(Dropout(0.3))
            model.add(Dense(32, activation='relu'))
            model.add(Dense(1, activation='sigmoid'))

            X_train_m = X_train.values.reshape((X_train.shape[0], 1, X_train.shape[1]))
            X_test_m  = X_test.values.reshape((X_test.shape[0], 1, X_test.shape[1]))

        elif model_name == "ANN":
            model = Sequential()
            model.add(tf.keras.Input(shape=(X_train.shape[1],)))
            model.add(Dense(128, activation='relu')); model.add(BatchNormalization()); model.add(Dropout(0.3))
            model.add(Dense(64, activation='relu')); model.add(BatchNormalization()); model.add(Dropout(0.3))
            model.add(Dense(32, activation='relu')); model.add(Dropout(0.3))
            model.add(Dense(1, activation='sigmoid'))
            X_train_m, X_test_m = X_train, X_test

        elif model_name == "GRU":
            model = Sequential()
            model.add(GRU(128, input_shape=(X_train.shape[1], 1), return_sequences=True))
            model.add(BatchNormalization()); model.add(Dropout(0.3))
            model.add(GRU(64, return_sequences=False))
            model.add(BatchNormalization()); model.add(Dropout(0.3))
            model.add(Dense(64, activation='relu')); model.add(Dropout(0.3))
            model.add(Dense(32, activation='relu'))
            model.add(Dense(1, activation='sigmoid'))
            X_train_m = X_train.values.reshape((X_train.shape[0], X_train.shape[1], 1))
            X_test_m  = X_test.values.reshape((X_test.shape[0], X_test.shape[1], 1))

        elif model_name == "CNN":
            model = Sequential()
            model.add(tf.keras.Input(shape=(X_train.shape[1], 1)))
            model.add(Conv1D(filters=64, kernel_size=min(2, X_train.shape[1]), activation='relu', padding='same'))
            model.add(BatchNormalization())
            if X_train.shape[1] >= 4:
                model.add(MaxPooling1D(pool_size=2))
            model.add(Dropout(0.3))
            model.add(Conv1D(filters=128, kernel_size=min(2, X_train.shape[1]), activation='relu', padding='same'))
            model.add(BatchNormalization())
            if X_train.shape[1] >= 8:
                model.add(MaxPooling1D(pool_size=2))
            model.add(Dropout(0.3))
            model.add(Flatten())
            model.add(Dense(64, activation='relu')); model.add(Dropout(0.3))
            model.add(Dense(32, activation='relu'))
            model.add(Dense(1, activation='sigmoid'))

            X_train_m = X_train.values.reshape((X_train.shape[0], X_train.shape[1], 1))
            X_test_m  = X_test.values.reshape((X_test.shape[0], X_test.shape[1], 1))

        else:
            raise ValueError(f"Unknown model name: {model_name}")

        model.compile(loss='binary_crossentropy', optimizer='adam', metrics=['accuracy'])
        # Use EarlyStopping to avoid wasting epochs if validation stops improving
        es = EarlyStopping(monitor='val_loss', patience=3, restore_best_weights=True)

        print("Training...")
        history = model.fit(
            X_train_m, y_train,
            validation_data=(X_test_m, y_test),
            epochs=TRAING_EPOCHS,
            batch_size=BATCH_SIZE,
            callbacks=[es],
            verbose=VERBOSE
        )
        print("Training finished.")

        # Evaluate with minimal verbosity
        train_loss, train_acc = model.evaluate(X_train_m, y_train, verbose=0)
        test_loss, test_acc = model.evaluate(X_test_m, y_test, verbose=0)

        y_train_pred = (model.predict(X_train_m, verbose=0) > 0.5).astype(int)
        y_test_pred  = (model.predict(X_test_m, verbose=0) > 0.5).astype(int)

        train_f1 = f1_score(y_train, y_train_pred)
        test_f1  = f1_score(y_test, y_test_pred)

        results.append({
            "k": k,
            "train_acc": float(train_acc),
            "test_acc": float(test_acc),
            "train_f1": float(train_f1),
            "test_f1": float(test_f1)
        })

        print(f"Result k={k} -> train_acc={train_acc:.4f}, test_acc={test_acc:.4f}, test_f1={test_f1:.4f}")

    results_df = pd.DataFrame(results)
    best_row = results_df.loc[results_df['test_f1'].idxmax()]
    print(f"\n{model_name} BEST -> k={best_row['k']}, test_f1={best_row['test_f1']:.4f}")
    return results_df, best_row

# ---------------------------
# Report generation
# ---------------------------
def add_df_to_doc(document, df, title=None):
    if title:
        document.add_heading(title, level=2)
    # create table
    table = document.add_table(rows=1, cols=len(df.columns))
    try:
        table.style = DOCX_TABLE_STYLE
    except Exception:
        pass
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = str(row[col])
    try:
        table.autofit = True
    except Exception:
        pass

# ---------------------------
# Master runner
# ---------------------------
def run_all_models(X_fin_capped, y_encoded, feature_method="pca"):
    print(f"\n######## RUNNING ALL MODELS (fusion={feature_method}) ########")
    mi_scores, chi_scores, f_scores = calculate_Score(X_fin_capped, y_encoded)
    document = Document()
    document.add_heading(f"{feature_method.upper()}_FULL_MODEL_RESULTS", level=1)
    models = ["LSTM", "ANN", "GRU", "CNN"]
    for model in models:
        df, best = run_single_model(model, X_fin_capped, y_encoded, feature_method, mi_scores, chi_scores, f_scores)
        add_df_to_doc(document, df, title=f"{model} Results")
        document.add_paragraph(f"🏆 Best k = {best['k']} | Test F1 = {best['test_f1']:.4f}")
        document.add_page_break()
    filename = f"ALL_MODEL_RESULTS_{feature_method.upper()}.docx"
    document.save(filename)
    print(f"Saved report: {filename}")

# ---------------------------
# Main entry
# ---------------------------
def main(data_path):
    print("Loading dataset:", data_path)
    df = pd.read_csv(data_path)
    print("Dataset loaded. Shape:", df.shape)

    # detect target column
    target_col = find_target_column(df)
    print(f"Using target column: '{target_col}'")

    # Drop NaN / inf
    df.replace([np.inf, -np.inf], np.nan, inplace=True)
    df.dropna(inplace=True)
    print("After dropna shape:", df.shape)

    # Separate features and target
    y = df[target_col].astype(int)
    X = df.drop(columns=[target_col])

    # Keep only numeric columns (required by MI / chi2)
    X = X.select_dtypes(include=[np.number])
    print("Numeric features kept:", X.shape[1])

    # Scale features to [0,1] for stable NN training and fusion methods
    X = safe_minmax_scale(X)
    print("Features scaled using MinMaxScaler.")

    # convert y to numpy array for stratify checks
    y_np = np.array(y)

    # quick safety check: ensure binary target
    unique = np.unique(y_np)
    if not set(unique).issubset({0,1}):
        raise ValueError("Target must be binary 0/1. Found values: " + str(unique))

    # run for each fusion method
    for method in ['pca','ica','fa']:
        run_all_models(X, y_np, feature_method=method)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python train_ids.py <data_csv_path>")
        sys.exit(1)
    data_path = sys.argv[1]
    start_time = time.time()
    main(data_path)
    print("Total elapsed time: {:.1f}s".format(time.time() - start_time))
    print("Done.")