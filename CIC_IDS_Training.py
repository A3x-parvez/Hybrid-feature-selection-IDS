import os
import sys
import glob
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
import matplotlib.pyplot as plt
from sklearn.decomposition import PCA
from sklearn.decomposition import FastICA
from sklearn.preprocessing import MinMaxScaler
from sklearn.decomposition import FactorAnalysis
from sklearn.model_selection import train_test_split
from sklearn.feature_selection import mutual_info_classif, chi2, f_classif
from sklearn.metrics import f1_score, classification_report, confusion_matrix

import tensorflow as tf

gpus = tf.config.list_physical_devices('GPU')
if gpus:          # <-- If GPU exists, do memory growth
    try:
        for gpu in gpus:
            tf.config.experimental.set_memory_growth(gpu, True)
        print("✅ Memory growth enabled on GPU.")
    except RuntimeError as e:
        print("⚠️ Could not set memory growth:", e)
else:
    print("⚠️ No GPU found. Running on CPU.")

from tensorflow.keras.layers import Conv1D, MaxPooling1D, Flatten, GRU
from tensorflow.keras.models import Sequential
from tensorflow.keras.utils import to_categorical
from tensorflow.keras.layers import LSTM, Dense, Dropout, Bidirectional,BatchNormalization
from tensorflow.keras.callbacks import EarlyStopping


# GLOBAL CONFIGURATION

if len(sys.argv) < 2:
    print("Usage: python CIC_IDS_Training.py <data_path>")
    sys.exit(1)

FUSION_LIST = ['pca','ica','fa']
MODEL_LIST = ['LSTM','ANN','GRU','CNN']
RANDOM_STATE = 42
TRAING_EPOCHS = 50
BATCH_SIZE = 512
SPLIT_SIZE = 0.2
RANGE_LIMIT = 75  # Upper limit for k in feature selection SHOULD BE 65 IN FINAL RUN


data_path = sys.argv[1]

def load_data_from_directory(data_path):
    # df= pd.read_csv("CIC-IDS-2017_fin_capped.csv")
    df = pd.read_csv(data_path)
    print("Dataset loaded sucessfully")
    return df

def preprocess_data(df):
    # Drop rows with NaN or infinite values
    df.replace([np.inf, -np.inf], np.nan, inplace=True)
    df.dropna(inplace=True)
    
    print(f"Data shape after dropping NaN and infinite values: {df.shape}")

    #split features and labels
    X_fin_capped = df.drop(columns=["target"])
    y_encoded = df["target"]
    
    print("Class distribution:")
    print(y_encoded.value_counts())
    
    print('Preprocessing done.')
    return X_fin_capped, y_encoded

def split_sub_data(X_fin_capped, y_encoded, selected, test_size=0.2, random_state=42):

    print("Selected Features:", selected)

    print("New data created .")
    # 2. Create new dataframe with selected features
    X_selected = X_fin_capped[selected]
    print("Feature selection done.")

    # 3. Show shape
    print("Original shape:", X_fin_capped.shape)
    print("Selected shape:", X_selected.shape)
    
    # Split into train and test
    from sklearn.model_selection import train_test_split
    X_train, X_test, y_train, y_test = train_test_split(X_selected, y_encoded, test_size=test_size, random_state=random_state)
    print("Train test split done.")

    # Convert to numpy arrays (if not already)
    y_train = np.array(y_train)
    y_test  = np.array(y_test)

    # ----- TRAINING SET COUNTS -----
    train_total  = len(y_train)
    train_class0 = np.sum(y_train == 0)
    train_class1 = np.sum(y_train == 1)

    # ----- TESTING SET COUNTS -----
    test_total  = len(y_test)
    test_class0 = np.sum(y_test == 0)
    test_class1 = np.sum(y_test == 1)

    print("\n====== Class Distribution After Split ======")
    print(f"Training Total  : {train_total}")
    print(f"Training Class 0: {train_class0}")
    print(f"Training Class 1: {train_class1}")
    print("-------------------------------------------")
    print(f"Testing Total   : {test_total}")
    print(f"Testing Class 0 : {test_class0}")
    print(f"Testing Class 1 : {test_class1}")
    print("===========================================\n")

    return X_train, X_test, y_train, y_test


def calculate_Score(X, y):
    
    print("calculated mi score.")
    mi_scores = mutual_info_classif(X, y, random_state=42)
    print("Sucessfully calculated mi score.")

    print("calculated chi score.")
    X_chi = X.copy()
    X_chi[X_chi < 0] = 0
    chi_scores, _ = chi2(X_chi, y)
    print("Sucessfully calculated chi score.")
   
    print("calculated f score.")
    f_scores, _ = f_classif(X, y)
    print("Sucessfully calculated f score.")

    return mi_scores,chi_scores,f_scores


def hybrid_feature_selection_pca(mi_scores, chi_scores, f_scores, feature_names, k):
    """
    PCA-based hybrid feature selection using MI, Chi², and F-score fusion.

    Parameters:
        mi_scores, chi_scores, f_scores : array-like
            Feature importance scores.
        feature_names : list or pd.Index
            Feature names corresponding to the scores.
        k : int
            Number of top features to select.

    Returns:
        selected_features : list
            Names of the top-k selected features.
    """
    print("Starting PCA-based feature fusion...")

    # Stack scores into a matrix [n_features x 3]
    scores = np.vstack([mi_scores, chi_scores, f_scores]).T

    # Normalize scores (so all are in the same scale)
    scaler = MinMaxScaler()
    scores_scaled = scaler.fit_transform(scores)

    # Apply PCA to combine the 3 score types into one meta-score
    pca = PCA(n_components=1)
    meta_score = pca.fit_transform(scores_scaled).ravel()

    # Create a DataFrame for clarity
    feature_df = pd.DataFrame({
        'Feature': feature_names,
        'MI_Score': mi_scores,
        'Chi2_Score': chi_scores,
        'F_Score': f_scores,
        'Meta_Score': meta_score
    }).sort_values(by='Meta_Score', ascending=False)

    # Select top-k features
    selected_features = feature_df.head(k)['Feature'].tolist()

    print("✅ PCA fusion completed.")
    print("Top", k, "Selected Features:", selected_features)
    print("\nPCA Component Weights (importance of MI, Chi², F):")
    print(pd.Series(pca.components_[0], index=['MI', 'Chi2', 'F']))

    return selected_features, feature_df


def hybrid_feature_selection_ica(mi_scores, chi_scores, f_scores, feature_names, k):
    """
    Hybrid feature selection using Independent Component Analysis (ICA).
    """
    print("Starting ICA-based feature fusion...")

    # 1. Stack scores [n_features x 3]
    scores = np.vstack([mi_scores, chi_scores, f_scores]).T

    # 2. Normalize scores (Crucial for ICA to converge)
    scaler = MinMaxScaler()
    scores_scaled = scaler.fit_transform(scores)

    # 3. Apply FastICA
    # We assume there is 1 independent source signal: "True Feature Quality"
    ica = FastICA(n_components=1, random_state=42, whiten='unit-variance')
    
    # Get the independent component
    meta_score = ica.fit_transform(scores_scaled).ravel()

    # --- CRITICAL STEP FOR ICA: Check for Sign Flipping ---
    # ICA cannot determine the sign of the source. It might give 
    # good features negative scores. We check this by correlating 
    # the ICA output with the average of the original scores.
    
    simple_average = np.mean(scores_scaled, axis=1)
    correlation = np.corrcoef(meta_score, simple_average)[0, 1]
    
    if correlation < 0:
        print("🔄 ICA sign flip detected. Inverting scores to match reality...")
        meta_score = -meta_score
    # ------------------------------------------------------

    # 4. Create DataFrame
    feature_df = pd.DataFrame({
        'Feature': feature_names,
        'MI_Score': mi_scores,
        'Chi2_Score': chi_scores,
        'F_Score': f_scores,
        'ICA_Meta_Score': meta_score
    }).sort_values(by='ICA_Meta_Score', ascending=False)

    # 5. Select top-k features
    selected_features = feature_df.head(k)['Feature'].tolist()

    print("✅ ICA fusion completed.")
    print(f"Top {k} Selected Features: {selected_features}")

    return selected_features, feature_df


def hybrid_feature_selection_fa(mi_scores, chi_scores, f_scores, feature_names, k):
    """
    Hybrid feature selection using Factor Analysis (FA) fusion.
    """
    print("Starting Factor Analysis (FA) feature fusion...")

    # 1. Stack scores [n_features x 3]
    scores = np.vstack([mi_scores, chi_scores, f_scores]).T

    # 2. Normalize scores (Essential for FA)
    scaler = MinMaxScaler()
    scores_scaled = scaler.fit_transform(scores)

    # 3. Apply Factor Analysis
    # We want 1 latent factor (the "True Importance")
    fa = FactorAnalysis(n_components=1, random_state=42)
    
    # The output is the feature's score on the hidden factor
    meta_score = fa.fit_transform(scores_scaled).ravel()

    # 4. Create DataFrame
    feature_df = pd.DataFrame({
        'Feature': feature_names,
        'MI_Score': mi_scores,
        'Chi2_Score': chi_scores,
        'F_Score': f_scores,
        'Meta_Score': meta_score
    }).sort_values(by='Meta_Score', ascending=False)

    # 5. Select top-k
    selected_features = feature_df.head(k)['Feature'].tolist()

    print("✅ FA fusion completed.")
    print(f"Top {k} Selected Features: {selected_features}")
    
    # Optional: Check how much FA relied on each metric
    print("\nFactor Loadings (Correlation with MI, Chi2, F):")
    print(pd.Series(fa.components_[0], index=['MI', 'Chi2', 'F']))

    return selected_features, feature_df



def feature_selection(X_fin_capped, mi_scores, chi_scores, f_scores, k=20, method='pca'):
    # # Calculate individual scores
    # print("Calculating individual feature scores...")
    # mi_scores, chi_scores, f_scores = calculate_Score(X_fin_capped, y_encoded)
    # print("Sucessfully calculated individual feature scores.")
    
    # Hybrid selection
    if method == 'pca':
        selected_features, score_table = hybrid_feature_selection_pca(mi_scores, chi_scores, f_scores, feature_names=X_fin_capped.columns, k=k)
        print("Selected Features:", selected_features)
    elif method == 'ica':
        selected_features, score_table = hybrid_feature_selection_ica(mi_scores, chi_scores, f_scores,  feature_names=X_fin_capped.columns, k=k)
        print("Selected Features:", selected_features)
    elif method == 'fa':
        selected_features, score_table = hybrid_feature_selection_fa(mi_scores, chi_scores, f_scores,  feature_names=X_fin_capped.columns, k=k)
        print("Selected Features:", selected_features)
    else:
        raise ValueError("Invalid method. Choose from 'pca', 'ica', or 'fa'.")

    return selected_features, score_table


# ===========================================
#  MODULE: UNIVERSAL TRAINING PIPELINE
# ===========================================
def run_single_model(model_name, X_fin_capped, y_encoded, feature_method, mi_scores, chi_scores, f_scores):

    results = []
    print(f"\n============== {model_name} TRAINING STARTED ==============")

    
    for k in range(5, RANGE_LIMIT, 5):

        print(f"\n===== Testing k = {k} =====")
        print("Feature selection start:")

        selected, value_table = feature_selection(
            X_fin_capped=X_fin_capped, mi_scores=mi_scores, chi_scores=chi_scores, f_scores=f_scores, k=k, method=feature_method
        )

        # SPLIT DATA (YOUR FUNCTION)
        X_train, X_test, y_train, y_test = split_sub_data(
            X_fin_capped=X_fin_capped, y_encoded=y_encoded, selected=selected, test_size=SPLIT_SIZE, random_state=RANDOM_STATE
        )

        # ---------------------------------------------------------
        # MODEL BUILDING — NO CHANGE IN MODEL LOGIC
        # ---------------------------------------------------------

        if model_name == "LSTM":
            # Your exact LSTM model (unchanged)
            model = Sequential()
            model.add(LSTM(128, input_shape=(1, X_train.shape[1]), return_sequences=True))
            model.add(BatchNormalization())
            model.add(Dropout(0.3))
            model.add(LSTM(64, return_sequences=True))
            model.add(Dropout(0.3))
            model.add(LSTM(64, return_sequences=False))
            model.add(BatchNormalization())
            model.add(Dropout(0.3))
            model.add(Dense(64, activation='relu'))
            model.add(Dropout(0.3))
            model.add(Dense(32, activation='relu'))
            model.add(Dense(1, activation='sigmoid'))

            # reshape (unchanged)
            X_train_m = X_train.values.reshape((X_train.shape[0], 1, X_train.shape[1]))
            X_test_m  = X_test.values.reshape((X_test.shape[0], 1, X_test.shape[1]))

        elif model_name == "ANN":
            model = Sequential()
            model.add(tf.keras.Input(shape=(X_train.shape[1],)))
            model.add(Dense(128, activation='relu'))
            model.add(BatchNormalization()); 
            model.add(Dropout(0.3))
            model.add(Dense(64, activation='relu'))
            model.add(BatchNormalization()); 
            model.add(Dropout(0.3))
            model.add(Dense(32, activation='relu')); 
            model.add(Dropout(0.3))
            model.add(Dense(1, activation='sigmoid'))
            X_train_m, X_test_m = X_train, X_test

        elif model_name == "GRU":
            model = Sequential()
            model.add(GRU(128, input_shape=(X_train.shape[1], 1), return_sequences=True))
            model.add(BatchNormalization()); 
            model.add(Dropout(0.3))
            model.add(GRU(64, return_sequences=False))
            model.add(BatchNormalization()); 
            model.add(Dropout(0.3))
            model.add(Dense(64, activation='relu')); 
            model.add(Dropout(0.3))
            model.add(Dense(32, activation='relu'))
            model.add(Dense(1, activation='sigmoid'))
            X_train_m = X_train.values.reshape((X_train.shape[0], 1, X_train.shape[1]))
            X_test_m  = X_test.values.reshape((X_test.shape[0], 1, X_test.shape[1]))

        elif model_name == "CNN":
            model = Sequential()
            model.add(tf.keras.Input(shape=(X_train.shape[1], 1)))

            model.add(Conv1D(filters=64, kernel_size=min(2, X_train.shape[1]),
                             activation='relu', padding='same'))
            model.add(BatchNormalization())
            # if X_train.shape[1] > 4:
            if X_train.shape[1] >= 4:
                model.add(MaxPooling1D(pool_size=2))
            model.add(Dropout(0.3))

            model.add(Conv1D(filters=128, kernel_size=min(2, X_train.shape[1]),
                             activation='relu', padding='same'))
            model.add(BatchNormalization())
            if X_train.shape[1] > 8:
                model.add(MaxPooling1D(pool_size=2))
            model.add(Dropout(0.3))

            model.add(Flatten())
            model.add(Dense(64, activation='relu')); 
            model.add(Dropout(0.3))
            model.add(Dense(32, activation='relu'))
            model.add(Dense(1, activation='sigmoid'))

            X_train_m = X_train.values.reshape((X_train.shape[0], X_train.shape[1], 1))
            X_test_m  = X_test.values.reshape((X_test.shape[0], X_test.shape[1], 1))

        # Compile (unchanged)
        model.compile(loss='binary_crossentropy', optimizer='adam', metrics=['accuracy'])
        model.summary()

        # # GPU settings (unchanged)
        # tf.debugging.set_log_device_placement(True)
        # tf.config.set_soft_device_placement(True)
        # ✅ Allow fallback to CPU if needed
        tf.debugging.set_log_device_placement(True)
        tf.config.set_soft_device_placement(True)  # <-- changed to True
        
        # Optional: Enable memory growth
        gpus = tf.config.list_physical_devices('GPU')
        if gpus:
            try:
                for gpu in gpus:
                    tf.config.experimental.set_memory_growth(gpu, True)
            except RuntimeError as e:
                print("Memory growth error:", e)
        
        # # ✅ Train without manual device forcing — TF will auto use GPU
        # early_stop = tf.keras.callbacks.EarlyStopping(monitor='val_loss', patience=3, restore_best_weights=True)


        print("Training started...")
        history = model.fit(
            X_train_m, y_train,
            validation_data=(X_test_m, y_test),
            epochs=TRAING_EPOCHS,
            batch_size=BATCH_SIZE,
            verbose=1
        )
        print("Training done.")

        # Evaluation (unchanged)
        train_loss, train_acc = model.evaluate(X_train_m, y_train, verbose=0)
        test_loss, test_acc = model.evaluate(X_test_m, y_test, verbose=0)

        y_train_pred = (model.predict(X_train_m) > 0.5).astype(int)
        y_test_pred  = (model.predict(X_test_m) > 0.5).astype(int)

        train_f1 = f1_score(y_train, y_train_pred)
        test_f1  = f1_score(y_test, y_test_pred)

        results.append({
            "k": k,
            "train_acc": train_acc,
            "test_acc": test_acc,
            "train_f1": train_f1,
            "test_f1": test_f1
        })

        print(f"Train Acc = {train_acc:.4f}, Test Acc = {test_acc:.4f}")
        print(f"Train F1  = {train_f1:.4f}, Test F1  = {test_f1:.4f}")

    results_df = pd.DataFrame(results)
    best_row = results_df.loc[results_df['test_f1'].idxmax()]
    print(f"\n{model_name} : 🏆 Best k = {best_row['k']} | Test F1 = {best_row['test_f1']:.4f}")

    return results_df, best_row


# ===========================================
#  MASTER FUNCTION TO RUN ALL MODELS       
# ===========================================
def run_all_models(X_fin_capped, y_encoded, mi_scores, chi_scores, f_scores, feature_method="pca"):

    print(f"\n########## RUNNING ALL MODELS WITH {feature_method.upper()} FEATURE FUSION ##########")

    # mi_scores, chi_scores, f_scores = calculate_Score(X_fin_capped, y_encoded)

    document = Document()
    document.add_heading(f"{feature_method.upper()}_FULL MODEL RESULTS", level=1)

    models = ["LSTM", "ANN", "GRU", "CNN"]

    for model in models:

        # df, best = run_single_model(model, X_fin_capped, y_encoded, feature_method)
        df, best =  run_single_model(
            model_name=model, 
            X_fin_capped=X_fin_capped, 
            y_encoded=y_encoded, 
            feature_method=feature_method, 
            mi_scores=mi_scores, 
            chi_scores=chi_scores, 
            f_scores=f_scores
            )

        document.add_heading(f"{model} Results", level=2)

        # Create table with header row
        table = document.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light List Accent 1'  # Optional but highly recommended

        # Header row
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)

        # Add the data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                value = row[col]
                row_cells[i].text = str(value)

        # Add Best-K text
        document.add_paragraph(
            f"\n🏆 Best k = {best['k']} | Test F1 = {best['test_f1']:.4f}"
        )
        document.add_page_break()
    # ⭐ NEW: filename based on feature fusion method

    # filename = f"ALL_MODEL_RESULTS_{feature_method.upper()}.docx"
    # Create output directory if not exists
    output_dir = "model_results"
    os.makedirs(output_dir, exist_ok=True)

    # Build full file path
    filename = os.path.join(output_dir, f"ALL_MODEL_RESULTS_{feature_method.upper()}.docx")
    document.save(filename)

    print(f"\n📄 ALL MODEL RESULTS SAVED SUCCESSFULLY AS: {filename}")



if __name__ == "__main__":
    #Step 1: Load Data
    print("Data path received:", data_path)
    df = load_data_from_directory(data_path)

    #Step 2: Preprocess Data
    X_fin_capped, y_encoded = preprocess_data(df)

    #Step 3: Calculate Individual Feature Scores
    mi_scores, chi_scores, f_scores = calculate_Score(X_fin_capped, y_encoded)

    #Step 4: Run All Models with Different Feature Fusion Methods
    fusion_list = ['pca','ica','fa']

    for method in fusion_list:
        run_all_models(X_fin_capped, y_encoded, mi_scores=mi_scores, chi_scores=chi_scores, f_scores=f_scores, feature_method=method)  

    print("All processes completed successfully.")

